import os
import pandas as pd
from docx import Document

from io import BytesIO
import pandas as pd
from docx import Document
import os

def docx_table_to_df(docx_file):
    # Accept file path or file-like object
    if isinstance(docx_file, str):
        doc = Document(docx_file)
    else:
        doc = Document(docx_file)
    table = doc.tables[0]
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data[1:], columns=data[0])
    return df

def csv_to_df(csv_file):
    return pd.read_csv(csv_file)

def load_default_data():
    """Load default files from data/ directory."""
    data_dir = os.path.join(os.path.dirname(__file__), '..', 'data')
    doctor_df = docx_table_to_df(os.path.join(data_dir, 'Doctor_Charges.docx'))
    insurance_df = docx_table_to_df(os.path.join(data_dir, 'Medicaid_Insurance_Rates.docx'))
    patient_df = docx_table_to_df(os.path.join(data_dir, 'Patient_Disease_Assignments.docx'))
    return doctor_df, insurance_df, patient_df

def prepare_billing_summary(doctor_df, insurance_df, patient_df):
    # Ensure patient_df has 'Assigned Doctor' and 'Insurance Company'
    if 'Assigned Doctor' not in patient_df.columns:
        doctors = ['Doctor Kelvin Nkansa', 'Doctor Lord Gyasi'] * 3
        patient_df['Assigned Doctor'] = [doctors[i] for i in range(len(patient_df))]
    # Ensure 'Insurance Company' in both patient_df and insurance_df
    if 'Insurance Company' not in patient_df.columns:
        patient_df['Insurance Company'] = 'Medicaid'
    if 'Insurance Company' not in insurance_df.columns:
        insurance_df['Insurance Company'] = 'Medicaid'

    # Merge on disease, ICD, and insurance company
    merged = pd.merge(
        patient_df,
        doctor_df,
        left_on=['Disease', 'ICD Code'],
        right_on=['Disease Name', 'ICD Code'],
        how='left'
    )
    merged = pd.merge(
        merged,
        insurance_df,
        left_on=['Disease Name', 'ICD Code', 'Insurance Company'],
        right_on=['Disease Name', 'ICD Code', 'Insurance Company'],
        how='left'
    )

    # Calculate charges and patient responsibility
    def get_doctor_charge(row):
        if row['Assigned Doctor'] == 'Doctor Kelvin Nkansa':
            return float(row['Doctor A Rate ($)'])
        else:
            return float(row['Doctor B Rate ($)'])
    merged['Doctor Charge'] = merged.apply(get_doctor_charge, axis=1)
    merged['Insurance Rate'] = merged.filter(regex='Rate').iloc[:, -1].astype(float)
    merged['Patient Owes'] = merged['Doctor Charge'] - merged['Insurance Rate']

    summary_cols = [
        'Patient Name', 'Patient ID', 'Disease', 'ICD Code', 'Assigned Doctor', 'Insurance Company',
        'Doctor Charge', 'Insurance Rate', 'Patient Owes'
    ]
    summary = merged[summary_cols]
    return summary
